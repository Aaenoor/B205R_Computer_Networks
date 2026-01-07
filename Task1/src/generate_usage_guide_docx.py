"""
Generate DOCX version of Usage Guide
Converts the markdown usage guide to a professional Word document
Author: B205 Computer Networks Project
Date: January 2026
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_usage_guide_docx():
    """Create Word document version of the usage guide"""
    doc = Document()

    # Title Page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("B205 COMPUTER NETWORKS\n\n")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 0, 139)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Chat Application\nUsage Guide")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph("\n" * 2)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(
        "Gisma University of Applied Sciences\n"
        "Department of Computer and Data Sciences\n"
        "Autumn 2025"
    )
    run.font.size = Pt(12)

    doc.add_page_break()

    # Table of Contents
    doc.add_heading("Table of Contents", 0)
    toc_items = [
        "1. Quick Start",
        "2. Server Setup",
        "3. Client Setup",
        "4. User Interface Overview",
        "5. Sending Messages",
        "6. Private Messaging",
        "7. File Transfers",
        "8. Photo Sharing",
        "9. Managing Contacts",
        "10. Configuration Options",
        "11. Troubleshooting"
    ]
    for item in toc_items:
        doc.add_paragraph(item)

    doc.add_page_break()

    # Quick Start
    doc.add_heading("1. Quick Start", 1)

    doc.add_heading("Step 1: Start the Server", 2)
    doc.add_paragraph("Open Command Prompt and run:")
    p = doc.add_paragraph("cd \"d:\\E Drive\\claude\\Assignments\\15th\\15.3\\Task1\"")
    p.runs[0].font.name = 'Courier New'
    p = doc.add_paragraph("python src/server.py")
    p.runs[0].font.name = 'Courier New'

    doc.add_heading("Step 2: Start a Client", 2)
    doc.add_paragraph("Open a new Command Prompt and run:")
    p = doc.add_paragraph("python src/client.py")
    p.runs[0].font.name = 'Courier New'

    doc.add_heading("Step 3: Enter Username", 2)
    doc.add_paragraph("When prompted, enter a unique username (1-20 characters).")

    doc.add_heading("Step 4: Start Chatting!", 2)
    doc.add_paragraph("Type your message and press Enter or click Send.")

    # Server Setup
    doc.add_heading("2. Server Setup", 1)

    doc.add_heading("Starting the Server", 2)
    steps = [
        "Open Command Prompt or Terminal",
        "Navigate to the project directory",
        "Run: python src/server.py"
    ]
    for step in steps:
        p = doc.add_paragraph(style='List Number')
        p.add_run(step)

    doc.add_heading("Server Output", 2)
    doc.add_paragraph("You should see:")
    doc.add_paragraph(
        "==================================================\n"
        "  B205 Computer Networks - Chat Server\n"
        "  Gisma University of Applied Sciences\n"
        "==================================================\n\n"
        "[SERVER] Chat server running on localhost:5000\n"
        "[SERVER] Waiting for connections..."
    )

    doc.add_heading("Server Events", 2)
    doc.add_paragraph("The server will display:")
    events = [
        "[SERVER] New connection from (ip, port) - When a client connects",
        "[SERVER] User 'username' connected - When user joins",
        "[CHAT] username: message - When messages are sent",
        "[PRIVATE] sender -> recipient: message - Private messages",
        "[FILE] username sent filename (size bytes) - File transfers",
        "[SERVER] User 'username' disconnected - When user leaves"
    ]
    for event in events:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(event)

    # Client Setup
    doc.add_heading("3. Client Setup", 1)

    doc.add_heading("Starting the Client", 2)
    steps = [
        "Open a new Command Prompt (keep server running)",
        "Navigate to the project directory",
        "Run: python src/client.py"
    ]
    for step in steps:
        p = doc.add_paragraph(style='List Number')
        p.add_run(step)

    doc.add_heading("Login Dialog", 2)
    doc.add_paragraph("A dialog will appear asking for:")
    login_fields = [
        "Server Host: The server's IP address (default: localhost)",
        "Server Port: The server's port (default: 5000)",
        "Username: Your display name (1-20 characters)"
    ]
    for field in login_fields:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(field)

    doc.add_heading("Connection Status", 2)
    doc.add_paragraph("- Green status bar: Connected successfully")
    doc.add_paragraph("- Red status bar: Disconnected or error")

    # User Interface Overview
    doc.add_heading("4. User Interface Overview", 1)

    doc.add_paragraph("The main interface consists of:")

    components = [
        ("Title Bar", "Shows application name and your username"),
        ("Chat Area", "Displays all messages with timestamps"),
        ("Online Users Panel", "Lists currently connected users"),
        ("Message Input", "Text field for typing messages"),
        ("Send Button", "Sends the typed message"),
        ("Photo Button", "Opens file picker for images"),
        ("File Button", "Opens file picker for documents"),
        ("Private Checkbox", "Enable private messaging mode"),
        ("Status Bar", "Shows connection status")
    ]

    table = doc.add_table(rows=len(components)+1, cols=2)
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Description'

    for i, (component, description) in enumerate(components, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = component
        row_cells[1].text = description

    # Sending Messages
    doc.add_heading("5. Sending Messages", 1)

    doc.add_heading("Broadcast Messages (To Everyone)", 2)
    steps = [
        "Type your message in the input field",
        "Press Enter or click Send",
        "Message appears to all connected users"
    ]
    for step in steps:
        p = doc.add_paragraph(style='List Number')
        p.add_run(step)

    doc.add_heading("Message Display Format", 2)
    doc.add_paragraph("[HH:MM] username: message content")

    doc.add_heading("System Messages", 2)
    doc.add_paragraph("System notifications appear differently:")
    doc.add_paragraph("=== username has joined the chat ===")
    doc.add_paragraph("=== username has left the chat ===")

    # Private Messaging
    doc.add_heading("6. Private Messaging", 1)

    doc.add_heading("Sending Private Messages", 2)
    steps = [
        "Select Recipient: Click a username in the Online Users panel",
        "Enable Private Mode: Check the Private checkbox",
        "Type Message: Enter your message",
        "Send: Press Enter or click Send"
    ]
    for step in steps:
        p = doc.add_paragraph(style='List Number')
        p.add_run(step)

    doc.add_heading("Private Message Display", 2)
    doc.add_paragraph("[HH:MM] [PRIVATE] sender -> recipient: message")

    doc.add_heading("Notes", 2)
    notes = [
        "Only sender and recipient can see private messages",
        "Private messages are stored in the database",
        "Uncheck Private to return to broadcast mode"
    ]
    for note in notes:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(note)

    # File Transfers
    doc.add_heading("7. File Transfers", 1)

    doc.add_heading("Supported File Types", 2)
    types = [
        "Documents: .txt, .pdf, .doc, .docx",
        "Archives: .zip",
        "Images: .png, .jpg, .jpeg, .gif, .bmp, .webp"
    ]
    for file_type in types:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(file_type)

    doc.add_heading("Maximum File Size", 2)
    doc.add_paragraph("Default: 10 MB (configurable in config.ini)")

    doc.add_heading("Sending Files", 2)
    steps = [
        "Click the File button",
        "Select a file from the file picker",
        "Wait for the upload to complete",
        "Recipients receive download notification"
    ]
    for step in steps:
        p = doc.add_paragraph(style='List Number')
        p.add_run(step)

    doc.add_heading("Receiving Files", 2)
    doc.add_paragraph("When someone sends a file:")
    steps = [
        "A notification appears: [12:05] alice sent file: document.pdf (125 KB)",
        "File is automatically saved to your Downloads folder",
        "Location: ~/Downloads/ChatApp/filename"
    ]
    for step in steps:
        p = doc.add_paragraph(style='List Number')
        p.add_run(step)

    # Photo Sharing
    doc.add_heading("8. Photo Sharing", 1)

    doc.add_heading("Supported Image Formats", 2)
    formats = ["PNG (.png)", "JPEG (.jpg, .jpeg)", "GIF (.gif)", "BMP (.bmp)", "WebP (.webp)"]
    for fmt in formats:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(fmt)

    doc.add_heading("Sending Photos", 2)
    steps = [
        "Click the Photo button",
        "Select an image file",
        "Photo is uploaded and displayed inline"
    ]
    for step in steps:
        p = doc.add_paragraph(style='List Number')
        p.add_run(step)

    doc.add_heading("Photo Display", 2)
    doc.add_paragraph("Photos appear directly in the chat area as thumbnails:")
    doc.add_paragraph("[12:10] alice sent a photo: sunset.jpg")
    doc.add_paragraph("[IMAGE THUMBNAIL]")

    # Configuration Options
    doc.add_heading("9. Configuration Options", 1)
    doc.add_paragraph("Edit config.ini to customize the application:")

    doc.add_heading("Server Settings", 2)
    doc.add_paragraph(
        "[SERVER]\n"
        "host = localhost      # Server bind address\n"
        "port = 5000           # Server port\n"
        "max_connections = 10  # Max simultaneous clients\n"
        "buffer_size = 4096    # Socket buffer size"
    )

    doc.add_heading("Client Settings", 2)
    doc.add_paragraph(
        "[CLIENT]\n"
        "default_server = localhost\n"
        "default_port = 5000\n"
        "reconnect_attempts = 3\n"
        "reconnect_delay = 5"
    )

    doc.add_heading("File Transfer Settings", 2)
    doc.add_paragraph(
        "[FILE_TRANSFER]\n"
        "max_file_size = 10485760   # 10 MB in bytes\n"
        "chunk_size = 8192          # Transfer chunk size\n"
        "allowed_extensions = .txt,.pdf,.png,.jpg,.jpeg,.gif,.doc,.docx,.zip"
    )

    # Troubleshooting
    doc.add_heading("10. Troubleshooting", 1)

    doc.add_heading("Cannot Connect to Server", 2)
    doc.add_paragraph("Problem: Client shows 'Connection refused' error")
    doc.add_paragraph("Solutions:")
    solutions = [
        "Verify server is running",
        "Check host and port match in config",
        "Check firewall allows port 5000",
        "Try 127.0.0.1 instead of localhost"
    ]
    for solution in solutions:
        p = doc.add_paragraph(style='List Number')
        p.add_run(solution)

    doc.add_heading("Username Already Taken", 2)
    doc.add_paragraph("Problem: 'Username already taken' error")
    doc.add_paragraph("Solution: Choose a different username. Each connected user must have a unique name.")

    doc.add_heading("File Transfer Failed", 2)
    doc.add_paragraph("Problem: File won't send")
    doc.add_paragraph("Solutions:")
    solutions = [
        "Check file size (must be under 10 MB)",
        "Verify file extension is allowed",
        "Ensure stable network connection"
    ]
    for solution in solutions:
        p = doc.add_paragraph(style='List Number')
        p.add_run(solution)

    doc.add_heading("Messages Not Appearing", 2)
    doc.add_paragraph("Problem: Sent messages don't show")
    doc.add_paragraph("Solutions:")
    solutions = [
        "Check connection status in status bar",
        "Restart client and reconnect",
        "Check server console for errors"
    ]
    for solution in solutions:
        p = doc.add_paragraph(style='List Number')
        p.add_run(solution)

    # Save document
    output_file = "../docs/USAGE_GUIDE.docx"
    doc.save(output_file)
    print(f"[SUCCESS] Usage guide saved to: {output_file}")
    return output_file

if __name__ == "__main__":
    print("Generating Usage Guide DOCX...")
    create_usage_guide_docx()
    print("Done!")
